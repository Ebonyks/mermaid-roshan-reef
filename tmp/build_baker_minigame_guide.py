from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
DOCX = OUT / "baker_minigame_phase_guide.docx"
PNG = OUT / "baker_minigame_visual_guide.png"

NAVY = "24364B"; TEAL = "2B7A78"; CORAL = "E8755B"; CREAM = "FFF8E8"; GOLD = "F1B844"; PLUM = "75507B"; INK = "23313D"; PALE = "EAF5F4"

def font(size, bold=False):
    candidates = ["C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/calibri.ttf"]
    for p in candidates:
        if Path(p).exists(): return ImageFont.truetype(p, size)
    return ImageFont.load_default()

def visual():
    im = Image.new("RGB", (1800, 1080), "#" + CREAM)
    d = ImageDraw.Draw(im)
    title = font(58, True); sub = font(28); h = font(32, True); body = font(23); small = font(20)
    d.text((90, 52), "THE GREAT CAKE SHOW", font=title, fill="#"+NAVY)
    d.text((92, 128), "Pastry Chef minigame - player flow and feedback loop", font=sub, fill="#"+TEAL)
    phases = [
        ("1", "SIFT", "Rub side to side", "Flour falls\ninto bowl", "40 px travel", TEAL),
        ("2", "POUR", "Press and hold", "Milk rises to\nthe line", "4.2 seconds", CORAL),
        ("3", "STIR", "Draw wide circles", "Batter spins\nand thickens", "5 turns", PLUM),
        ("4", "BAKE", "Wait, then tap", "Cake rises\nand turns gold", "12 sec gold", GOLD),
        ("5", "PIPE", "Trace dotted ring", "Dots become\nfrosting beads", "14 dots", TEAL),
        ("6", "DECORATE", "Swim close + tap", "Cherries pop\nonto cake", "4 toppings", CORAL),
    ]
    x0, y, cardw, gap = 80, 270, 245, 38
    for i, (num, name, gesture, feedback, measure, col) in enumerate(phases):
        x = x0 + i*(cardw+gap)
        d.rounded_rectangle((x,y,x+cardw,y+525), radius=28, fill="white", outline="#"+col, width=6)
        d.ellipse((x+88,y+30,x+177,y+119), fill="#"+col)
        tw = d.textbbox((0,0),num,font=font(40,True))[2]
        d.text((x+132-tw/2,y+49), num, font=font(40,True), fill="white")
        tw = d.textbbox((0,0),name,font=h)[2]
        d.text((x+132-tw/2,y+148), name, font=h, fill="#"+NAVY)
        d.line((x+42,y+202,x+222,y+202), fill="#"+col,width=4)
        for line_i, text in enumerate(gesture.split("\n")):
            tw = d.textbbox((0,0),text,font=body)[2]
            d.text((x+132-tw/2,y+230+line_i*30), text, font=body, fill="#"+INK)
        for line_i, text in enumerate(feedback.split("\n")):
            tw = d.textbbox((0,0),text,font=body)[2]
            d.text((x+132-tw/2,y+320+line_i*30), text, font=body, fill="#"+TEAL)
        d.rounded_rectangle((x+32,y+432,x+233,y+486),radius=15,fill="#"+PALE)
        tw = d.textbbox((0,0),measure,font=small)[2]
        d.text((x+132-tw/2,y+448),measure,font=small,fill="#"+NAVY)
        if i < len(phases)-1:
            ax = x + cardw + 7
            d.line((ax, y+510, ax+18, y+510), fill="#"+GOLD, width=6)
            d.polygon([(ax+24,y+510),(ax+14,y+500),(ax+14,y+520)],fill="#"+GOLD)
    d.rounded_rectangle((90, 865, 1710, 1010), radius=22, fill="#"+NAVY)
    d.text((125,890), "SAFETY NET: early oven taps say 'Not yet'; at 34 seconds the cake advances automatically.", font=font(23), fill="white")
    d.text((125,932), "Progress is always voiced, pictured, and reinforced with sparkles and chimes.", font=font(23), fill="white")
    im.save(PNG)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tcPr.append(shd)

def cell_margin(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for side, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn('w:'+side))
        if node is None: node = OxmlElement('w:'+side); tcMar.append(node)
        node.set(qn('w:w'), str(value)); node.set(qn('w:type'),'dxa')

def set_cell(cell, text, bold=False, color=INK, size=9.5):
    cell.text = ""; p = cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08
    r = p.add_run(text); r.bold=bold; r.font.name="Calibri"; r._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); r._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell_margin(cell)

def fixed_table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
    tPr=t._tbl.tblPr; tblW=OxmlElement('w:tblW'); tblW.set(qn('w:w'),'9360'); tblW.set(qn('w:type'),'dxa'); tPr.append(tblW)
    grid=t._tbl.tblGrid
    for gridcol,w in zip(grid.gridCol_lst,widths): gridcol.set(qn('w:w'),str(w))
    for c,h,w in zip(t.rows[0].cells,headers,widths): c.width=Inches(w/1440); shade(c,"DCEDEA"); set_cell(c,h,True,NAVY,9)
    for row in rows:
        cells=t.add_row().cells
        for c,v,w in zip(cells,row,widths): c.width=Inches(w/1440); set_cell(c,v)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def p(doc,text="",style=None, bold_prefix=None):
    para=doc.add_paragraph(style=style)
    para.paragraph_format.space_after=Pt(6); para.paragraph_format.line_spacing=1.25
    if bold_prefix and text.startswith(bold_prefix):
        r=para.add_run(bold_prefix); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); para.add_run(text[len(bold_prefix):])
    else: para.add_run(text)
    return para

def setup(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.8); sec.right_margin=Inches(0.8)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name,size,col,before,after in [('Heading 1',16,TEAL,16,8),('Heading 2',13,NAVY,11,6),('Heading 3',11.5,PLUM,8,4)]:
        s=styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=footer.add_run('Mermaid Roshan: Reef of Light | Baker minigame guide'); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string('6B7785')

def bullet(doc,text):
    x=doc.add_paragraph(style='List Bullet'); x.paragraph_format.space_after=Pt(3); x.paragraph_format.line_spacing=1.18; x.add_run(text)

def main():
    OUT.mkdir(exist_ok=True); visual()
    doc=Document(); setup(doc)
    title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.paragraph_format.space_before=Pt(28); title.paragraph_format.space_after=Pt(7)
    r=title.add_run('THE GREAT CAKE SHOW'); r.bold=True; r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); r.font.size=Pt(29); r.font.color.rgb=RGBColor.from_string(NAVY)
    st=doc.add_paragraph(); st.alignment=WD_ALIGN_PARAGRAPH.CENTER; st.paragraph_format.space_after=Pt(18); rr=st.add_run('Pastry Chef / “Baker” minigame - implementation breakdown and visual guide'); rr.italic=True; rr.font.size=Pt(13); rr.font.color.rgb=RGBColor.from_string(TEAL)
    doc.add_picture(str(PNG),width=Inches(6.85)); cap=doc.add_paragraph('Visual guide - six distinct gestures form one forgiving cake-making performance.'); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic=True; cap.runs[0].font.size=Pt(9); cap.paragraph_format.space_after=Pt(12)
    p(doc,'Scope: current implementation in scripts/opera_house.gd and scripts/opera_act.gd, reviewed 2026-08-01. “Baker” is the Pastry Chef act, The Great Cake Show, on Floor 1 of the Pearl Opera House.')
    doc.add_heading('At a glance',1)
    fixed_table(doc,['Stage','Player gesture','Completion / visible response'],[
        ('Prelude','Swim + sparkle taps','Free the farmers from 6 mischief imps; their carrots become an optional carrot-cake variant.'),
        ('1. Sift','Side-to-side rub','40 px horizontal travel; flour flakes fall into the bowl.'),
        ('2. Pour','Press and hold','4.2 seconds; jug tips and milk rises to the orange line.'),
        ('3. Stir','Wide circular drag','Each full circle adds a stir, speeds the bowl, and plays a rising chime; current threshold is 5 turns.'),
        ('4. Bake','Wait, then tap','Cake rises for 12 seconds, turns golden, then accepts the tap.'),
        ('5. Pipe','Trace dotted ring','Pass over 14 dots; each becomes a frosting bead.'),
        ('6. Decorate','Swim close + tap','Tap four pink spots; a cherry pops onto each one, then the act wins.')], [1180,2100,6080])
    doc.add_heading('Experience contract',1)
    for x in ['Designed for a non-reader: every stage has a voiced prompt, a picture-led target, HUD status, a pulsing pointer after inactivity, sparkles, and chimes.', 'No fail state: a wrong/early oven tap only says “Not yet”; a distracted cook advances automatically after 34 seconds in the oven.', 'Input varies deliberately: scrub, hold, circular drag, timed tap, trace, then proximity tap. It avoids repeating one gesture across the recipe.', 'The stage is a pastry kitchen: reef backdrop, warm oven alcove at stage left, ingredient shelf at right, bowl upstage centre, and Pastry Chef GLB art with primitive fallbacks.'] : bullet(doc,x)
    doc.add_page_break(); doc.add_heading('Phase-by-phase breakdown',1)
    phases=[
        ('Prelude - Rescue and ingredient story','The act is configured as a “shell” performance with six imps and a farmers rescue. Roshan pops the imps before the curtain opens. If the farmers’ carrot gift is in the persistent opera pantry, three carrots appear beside the bowl and the opening voice identifies the result as carrot cake. This changes the story and visuals, not the core input sequence.','Transition: _build_order() immediately calls _begin_sift() for the chef finale. The legacy picture-order layer pads remain built as shared-engine scenery, but are not part of the live cake sequence.'),
        ('1. Sift - flour into the bowl','The game enables drag mode and places the sieve over the centre bowl. The child rubs a finger horizontally; only sideways distance counts. Every roughly 2 progress points creates a falling flour flake, so the gesture is visibly legible rather than a hidden meter.','Target: 40 units of accumulated horizontal movement (SIFT_NEED). Releasing the finger is harmless; it simply resets the “previous drag” position. Completion transitions directly to pouring.'),
        ('2. Pour - hold the milk jug','A white jug appears above-left of the bowl, with an orange fill line as the visual goal. While the finger is down, the jug tilts toward -62 degrees, sparkles appear periodically, and the milk mesh rises. Releasing lets the jug settle back but does not erase progress.','Target: 4.2 seconds held (POUR_NEED). The milk display permits a little visual overfill, but the stage moves on at the threshold. Drag control is cleared and the child is asked to stir.'),
        ('3. Stir - circles around the bowl','Roshan must be within 7 world units of the bowl. At that point movement hands over to finger drag. A drag must begin at least 40 screen pixels from the bowl centre, then its accumulated angular travel is measured. Each full revolution adds one stir, rotates the bowl, intensifies the chime, and starts the bowl’s active art state on the first turn.','Target: 5 completed revolutions in the current code. After 26 seconds at the bowl, a gentle rescue completes one stir automatically. Note: the HUD currently labels this as “/ 3”; that label is inconsistent with the actual 5-turn threshold and is worth correcting separately.'),
        ('4. Bake - watch for golden','The batter becomes a separate cake object inside the glowing oven alcove. It starts squat, rises smoothly during the first 12 seconds, and then switches from pale to gold with a sparkle burst and voice prompt. The pointer targets the oven/cake.','Target: tap after the cake is golden (12 seconds). An early tap gives a pale-cake sparkle and “Not yet” hint, with no reset. If no action happens by 34 seconds, the game calls the same success path automatically.'),
        ('5. Pipe - trace the frosting ring','Drag mode returns. Fourteen pink dots form a circular guide around the cake. Passing the finger within 62 screen pixels of any visible dot hides it and creates a larger frosting bead with a pop tween. The dots may be collected in any order, keeping the trace forgiving.','Target: all 14 guide dots. The last dot counts even if the finger lifts on that frame; drag mode then turns off and the decoration targets open.'),
        ('6. Decorate - add cherries and finish','Four pink topping targets appear in an arc around the bowl/cake. Roshan swims near a target and taps the action control; the target vanishes and a red cherry scales up onto the cake. Pitch rises with each chime and sparkles mark each successful placement.','Target: 4 toppings from the chef configuration. Completing the final target calls _win(), which runs the celebration and the configured carrot-cake win line.')]
    for name,detail,gate in phases:
        doc.add_heading(name,2); p(doc,detail); p(doc,'Implementation gate: '+gate, bold_prefix='Implementation gate: ')
    doc.add_heading('Current implementation notes',1)
    fixed_table(doc,['Topic','What is currently shipped','Why it matters for review'],[
        ('Recipe order','The chef begins at sift; it does not ask the child to deliver the vanilla/coral/plum layers, despite older design notes and shared “order” scaffolding.','Test the live progression, not the historical spec.'),
        ('Stir count','Code finishes at 5 full circular accumulations; the HUD displays “/ 3.”','This is the only clear player-facing count mismatch found in the phase flow.'),
        ('Decoration gesture','The published redesign calls for drag-and-drop cherries. The current handler activates a spot when Roshan is within 4.5 world units and the action button is tapped.','The guide records actual behavior; decide separately whether to change implementation or documentation.'),
        ('Baking timer','The design notes describe late taps as “still fine, just fewer sparkles.” Current code has no late-sparkle grading; it auto-advances at 34 seconds.','Useful distinction for future polish or probe expectations.')], [1700,4600,3060])
    doc.add_heading('QA walkthrough',1)
    for x in ['Start a new Pastry Chef show; confirm the rescue opens the stage and carrots appear only when the farmers’ gift is owned.', 'Perform each gesture once as named: horizontal scrub, held press, wide circular drag, golden-timing tap, ring trace, then proximity tap.', 'Test the safety net: tap the oven before gold, release during pouring, and wait past 34 seconds in bake. None should restart or lose progress.', 'Confirm feedback remains readable on the target tablet: voice line, target/pointer, HUD meter or count, visible transformation, and chime/sparkle response.'] : bullet(doc,x)
    doc.core_properties.title='The Great Cake Show - Baker Minigame Phase Guide'; doc.core_properties.subject='Implementation review and visual guide'; doc.core_properties.author='Codex'
    doc.save(DOCX)

if __name__ == '__main__': main()
